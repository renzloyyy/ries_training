from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT_PATH = "/Users/joeelysiummamolo/Downloads/ries_training/Research_Program_Overview_Script.docx"


SCRIPT_TITLE = "Research Program Dashboard Overview Presentation Script"

SCRIPT_PARAGRAPHS = [
    "Good day everyone. This dashboard presents the overall performance of the research program by showing the flow from proposals, to completed research, to publications, and to fund allocation.",
    "We begin with the headline indicators at the top. These cards summarize the scale of research activity: total proposals submitted, total completed research, total published papers, and total fund allocation. These metrics immediately show the size of the research program and its current output.",
    "Next, the trend chart shows total research activity over time. This helps us see whether research submissions, completed outputs, and published papers are increasing or declining across the selected years. From this chart, we can identify growth periods, slowdowns, and the general direction of institutional research productivity.",
    "Below that, we look at campus contribution. The campus output card shows which campus contributes the most research outputs. The completion rate by campus compares how effectively each campus converts proposals into completed work or outputs. The funding share by campus shows how resources are distributed across campuses.",
    "Together, these visuals tell a simple story: how much research enters the pipeline, how much is completed, how much becomes publication output, and which campuses are driving performance.",
    "The main value of this overview is that it gives decision-makers a quick understanding of research scale, progress, and contribution before moving into the more detailed proposal, publication, and funding tabs.",
]


def set_page_margins(section):
    # Keep the page layout clean and readable with standard report margins.
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)


def style_title(paragraph):
    # Apply explicit title styling instead of relying on Word defaults.
    run = paragraph.add_run(SCRIPT_TITLE)
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(6, 14, 64)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(10)


def add_bottom_rule(paragraph):
    # Add a subtle divider under the title to separate the heading from the script body.
    paragraph_properties = paragraph._p.get_or_add_pPr()
    border = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "6")
    bottom.set(qn("w:color"), "D9E4FF")
    border.append(bottom)
    paragraph_properties.append(border)


def add_body_paragraph(document, text):
    # Body copy uses a moderate size so the script is easy to read while presenting.
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(10)
    paragraph.paragraph_format.line_spacing = 1.25
    run = paragraph.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(33, 43, 66)


def build_document():
    document = Document()
    set_page_margins(document.sections[0])

    title = document.add_paragraph()
    style_title(title)
    add_bottom_rule(title)

    subtitle = document.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(14)
    subtitle_run = subtitle.add_run("Overview speaking script for the dashboard presentation")
    subtitle_run.font.name = "Arial"
    subtitle_run.font.size = Pt(11)
    subtitle_run.italic = True
    subtitle_run.font.color.rgb = RGBColor(90, 100, 120)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for paragraph_text in SCRIPT_PARAGRAPHS:
        add_body_paragraph(document, paragraph_text)

    document.save(OUTPUT_PATH)


if __name__ == "__main__":
    build_document()
